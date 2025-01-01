import logging
from telegram import Update, BotCommand
from telegram.ext import (
    ApplicationBuilder, CommandHandler, MessageHandler, ConversationHandler, filters, ContextTypes
)
from docx import Document

# Настройка логирования
logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.DEBUG
)
logger = logging.getLogger(__name__)

# Пути к файлам
TEMPLATE_PATH = r"C:\\Users\\Admin\\Desktop\\123.docx"
OUTPUT_PATH = r"C:\\Users\\Admin\\Desktop\\filled_dkp.docx"

# Этапы заполнения
(
    PLACE, DATE, SELLER_NAME, SELLER_BIRTHDATE, SELLER_ADDRESS, SELLER_PASSPORT, BUYER_NAME,
    BUYER_BIRTHDATE, BUYER_ADDRESS, BUYER_PASSPORT, VEHICLE_BRAND, VEHICLE_CATEGORY,
    VEHICLE_TYPE, REG_SIGN, VIN, YEAR, ENGINE, CHASSIS, BODY, COLOR, PTS_INFO, PRICE,
    REG_CERTIFICATE_INFO
) = range(23)

# Замена текста в шаблоне Word
def replace_text_in_paragraphs(paragraphs, replacements):
    for paragraph in paragraphs:
        for run in paragraph.runs:
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)

def replace_text_in_tables(tables, replacements):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, replacements)

# Функция для обработки пропуска этапов
async def skip_stage(update: Update, context: ContextTypes.DEFAULT_TYPE):
    current_stage = context.user_data.get('current_stage', PLACE)
    context.user_data[current_stage] = "Не указано"
    next_stage = current_stage + 1
    context.user_data['current_stage'] = next_stage
    return await stages[next_stage](update, context)

# Этапы заполнения договора
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data.clear()
    context.user_data['current_stage'] = PLACE
    await update.message.reply_text("Привет! Где заключается договор?")
    return PLACE

async def ask_date(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['place'] = update.message.text.strip()
    context.user_data['current_stage'] = DATE
    await update.message.reply_text("Укажите дату заключения договора (например, 01.01.2024):")
    return DATE

async def ask_seller_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['date'] = update.message.text.strip()
    context.user_data['current_stage'] = SELLER_NAME
    await update.message.reply_text("Введите ФИО продавца:")
    return SELLER_NAME

async def ask_seller_birthdate(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['seller_name'] = update.message.text.strip()
    context.user_data['current_stage'] = SELLER_BIRTHDATE
    await update.message.reply_text("Введите дату рождения продавца:")
    return SELLER_BIRTHDATE

async def ask_seller_address(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['seller_birthdate'] = update.message.text.strip()
    context.user_data['current_stage'] = SELLER_ADDRESS
    await update.message.reply_text("Введите адрес регистрации продавца:")
    return SELLER_ADDRESS

async def ask_seller_passport(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['seller_address'] = update.message.text.strip()
    context.user_data['current_stage'] = SELLER_PASSPORT
    await update.message.reply_text("Введите паспортные данные продавца (серия, номер, кем выдан):")
    return SELLER_PASSPORT

# Этапы для покупателя
async def ask_buyer_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['seller_passport'] = update.message.text.strip()
    context.user_data['current_stage'] = BUYER_NAME
    await update.message.reply_text("Введите ФИО покупателя:")
    return BUYER_NAME

async def ask_buyer_birthdate(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['buyer_name'] = update.message.text.strip()
    context.user_data['current_stage'] = BUYER_BIRTHDATE
    await update.message.reply_text("Введите дату рождения покупателя:")
    return BUYER_BIRTHDATE

async def ask_buyer_address(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['buyer_birthdate'] = update.message.text.strip()
    context.user_data['current_stage'] = BUYER_ADDRESS
    await update.message.reply_text("Введите адрес регистрации покупателя:")
    return BUYER_ADDRESS

async def ask_buyer_passport(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['buyer_address'] = update.message.text.strip()
    context.user_data['current_stage'] = BUYER_PASSPORT
    await update.message.reply_text("Введите паспортные данные покупателя (серия, номер, кем выдан):")
    return BUYER_PASSPORT

# Этапы для транспортного средства
async def ask_vehicle_brand(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['buyer_passport'] = update.message.text.strip()
    context.user_data['current_stage'] = VEHICLE_BRAND
    await update.message.reply_text("Введите марку и модель транспортного средства:")
    return VEHICLE_BRAND

async def ask_vehicle_category(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['vehicle_brand'] = update.message.text.strip()
    context.user_data['current_stage'] = VEHICLE_CATEGORY
    await update.message.reply_text("Введите категорию транспортного средства:")
    return VEHICLE_CATEGORY

async def ask_vehicle_type(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['vehicle_category'] = update.message.text.strip()
    context.user_data['current_stage'] = VEHICLE_TYPE
    await update.message.reply_text("Введите тип транспортного средства по ПТС:")
    return VEHICLE_TYPE

async def ask_reg_sign(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['vehicle_type'] = update.message.text.strip()
    context.user_data['current_stage'] = REG_SIGN
    await update.message.reply_text("Введите регистрационный знак:")
    return REG_SIGN

async def ask_vin(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['reg_sign'] = update.message.text.strip()
    context.user_data['current_stage'] = VIN
    await update.message.reply_text("Введите идентификационный номер (VIN):")
    return VIN

async def ask_year(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['vin'] = update.message.text.strip()
    context.user_data['current_stage'] = YEAR
    await update.message.reply_text("Введите год выпуска:")
    return YEAR

async def ask_engine(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['year'] = update.message.text.strip()
    context.user_data['current_stage'] = ENGINE
    await update.message.reply_text("Введите данные двигателя:")
    return ENGINE

async def ask_chassis(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['engine'] = update.message.text.strip()
    context.user_data['current_stage'] = CHASSIS
    await update.message.reply_text("Введите данные шасси:")
    return CHASSIS

async def ask_body(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['chassis'] = update.message.text.strip()
    context.user_data['current_stage'] = BODY
    await update.message.reply_text("Введите данные кузова:")
    return BODY

async def ask_color(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['body'] = update.message.text.strip()
    context.user_data['current_stage'] = COLOR
    await update.message.reply_text("Введите цвет транспортного средства:")
    return COLOR

async def ask_pts_info(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['color'] = update.message.text.strip()
    context.user_data['current_stage'] = PTS_INFO
    await update.message.reply_text("Введите данные ПТС (серия, номер, кем выдан):")
    return PTS_INFO

async def ask_price(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['pts_info'] = update.message.text.strip()
    context.user_data['current_stage'] = PRICE
    await update.message.reply_text("Введите стоимость транспортного средства (например, 1000000):")
    return PRICE

async def ask_reg_certificate_info(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['price'] = update.message.text.strip()
    context.user_data['current_stage'] = REG_CERTIFICATE_INFO
    await update.message.reply_text("Введите данные свидетельства о регистрации (серия, номер, кем выдан):")
    return REG_CERTIFICATE_INFO

async def generate_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['reg_certificate_info'] = update.message.text.strip()
    doc = Document(TEMPLATE_PATH)

    # Подготовка данных для замены
    replacements = {f"{{{{{key}}}}}": value for key, value in context.user_data.items() if key != 'current_stage'}
    logger.debug(f"Реплейсменты: {replacements}")

    # Замена меток в параграфах
    replace_text_in_paragraphs(doc.paragraphs, replacements)

    # Замена меток в таблицах (если есть таблицы)
    replace_text_in_tables(doc.tables, replacements)

    # Сохранение документа
    doc.save(OUTPUT_PATH)
    logger.info(f"Документ сохранён по адресу: {OUTPUT_PATH}")

    # Отправка документа пользователю
    try:
        with open(OUTPUT_PATH, 'rb') as file:
            await update.message.reply_document(
                document=file,
                filename="filled_dkp.docx",
                caption="Документ успешно создан и заполнен!"
            )
    except Exception as e:
        logger.error(f"Ошибка отправки документа: {e}")
        await update.message.reply_text("Ошибка при создании или отправке документа. Попробуйте ещё раз.")
    return ConversationHandler.END


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Процесс заполнения договора отменён.")
    return ConversationHandler.END

# Настройка обработчиков
stages = {
    PLACE: ask_date,
    DATE: ask_seller_name,
    SELLER_NAME: ask_seller_birthdate,
    SELLER_BIRTHDATE: ask_seller_address,
    SELLER_ADDRESS: ask_seller_passport,
    SELLER_PASSPORT: ask_buyer_name,
    BUYER_NAME: ask_buyer_birthdate,
    BUYER_BIRTHDATE: ask_buyer_address,
    BUYER_ADDRESS: ask_buyer_passport,
    BUYER_PASSPORT: ask_vehicle_brand,
    VEHICLE_BRAND: ask_vehicle_category,
    VEHICLE_CATEGORY: ask_vehicle_type,
    VEHICLE_TYPE: ask_reg_sign,
    REG_SIGN: ask_vin,
    VIN: ask_year,
    YEAR: ask_engine,
    ENGINE: ask_chassis,
    CHASSIS: ask_body,
    BODY: ask_color,
    COLOR: ask_pts_info,
    PTS_INFO: ask_price,
    PRICE: ask_reg_certificate_info,
    REG_CERTIFICATE_INFO: generate_document,
}

# Основная функция
def main():
    app = ApplicationBuilder().token("8038967691:AAFWzfUkivYoCppp7qip4xBT-r3dxQeZcxU").build()

    app.bot.set_my_commands([
        BotCommand("start", "Начать заполнение договора"),
        BotCommand("cancel", "Отменить процесс заполнения"),
        BotCommand("skip", "Пропустить текущий этап")
    ])

    conversation_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={stage: [MessageHandler(filters.TEXT & ~filters.COMMAND, func)] for stage, func in stages.items()},
        fallbacks=[
            CommandHandler("cancel", cancel),
            CommandHandler("skip", skip_stage)
        ]
    )

    app.add_handler(conversation_handler)
    app.run_polling()

if __name__ == "__main__":
    main()
